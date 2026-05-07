from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent.parent
THESIS_PATH = ROOT / "thesis" / "实验进行中-gpt修改版本.docx"
BACKUP_PATH = ROOT / "thesis" / "实验进行中-gpt修改版本.before-template-format-backup.docx"


def set_east_asia_font(style, font_name: str, size_pt: float | None = None, bold: bool | None = None) -> None:
    style.font.name = font_name
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def iter_paragraphs(doc: Document):
    return list(doc.paragraphs)


def find_paragraph(doc: Document, predicate) -> Paragraph:
    for paragraph in iter_paragraphs(doc):
        if predicate(paragraph):
            return paragraph
    raise ValueError("Unable to find matching paragraph")


def paragraph_text(paragraph: Paragraph) -> str:
    return paragraph.text.strip()


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def insert_paragraph_before(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    if style:
        new_paragraph.style = style
    return new_paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    if style:
        new_paragraph.style = style
    return new_paragraph


def has_sect_pr(paragraph: Paragraph) -> bool:
    p_pr = paragraph._p.pPr
    return p_pr is not None and p_pr.sectPr is not None


def clone_sect_pr(paragraph: Paragraph):
    p_pr = paragraph._p.pPr
    if p_pr is None or p_pr.sectPr is None:
        raise ValueError("Paragraph does not contain section properties")
    return deepcopy(p_pr.sectPr)


def add_sect_pr(paragraph: Paragraph, sect_pr) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:sectPr"))
    if existing is not None:
        p_pr.remove(existing)
    p_pr.append(deepcopy(sect_pr))


def remove_sect_pr(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return
    sect_pr = p_pr.find(qn("w:sectPr"))
    if sect_pr is not None:
        p_pr.remove(sect_pr)


def strip_header_footer_refs(sect_pr) -> None:
    for tag in ("w:headerReference", "w:footerReference"):
        for element in list(sect_pr.findall(qn(tag))):
            sect_pr.remove(element)


def set_page_numbering(section, fmt: str, start: int | None = None) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))
    sect_pr.append(pg_num_type)


def append_field(paragraph: Paragraph, instruction: str, placeholder: str = "") -> None:
    begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin._r.append(fld_begin)

    instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    instr._r.append(instr_text)

    separate = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    separate._r.append(fld_sep)

    if placeholder:
        paragraph.add_run(placeholder)

    end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end._r.append(fld_end)


def build_toc_field(paragraph: Paragraph) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    append_field(paragraph, 'TOC \\\\o "1-3" \\\\h \\\\z \\\\u', "右键更新目录")


def build_page_field(paragraph: Paragraph) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    append_field(paragraph, "PAGE", "1")


def build_header_text(paragraph: Paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(10.5)
    r_fonts = run._r.get_or_add_rPr().rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        run._r.get_or_add_rPr().append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "宋体")


def build_header_styleref(paragraph: Paragraph) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    append_field(paragraph, ' STYLEREF "Heading 1" ', "当前章标题")
    for run in paragraph.runs:
        run.font.name = "宋体"
        run.font.size = Pt(10.5)
        r_fonts = run._r.get_or_add_rPr().rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            run._r.get_or_add_rPr().append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), "宋体")


def clear_story(story) -> Paragraph:
    for paragraph in list(story.paragraphs)[1:]:
        delete_paragraph(paragraph)
    first = story.paragraphs[0]
    clear_paragraph(first)
    return first


def normalize_styles(doc: Document) -> None:
    title = doc.styles["Title"]
    set_east_asia_font(title, "黑体", 18, True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading1 = doc.styles["Heading 1"]
    set_east_asia_font(heading1, "黑体", 16, True)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading2 = doc.styles["Heading 2"]
    set_east_asia_font(heading2, "黑体", 15, True)
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    heading3 = doc.styles["Heading 3"]
    set_east_asia_font(heading3, "黑体", 14, True)
    heading3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    style_names = {style.name for style in doc.styles}
    if "Caption" not in style_names:
        doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)

    caption = doc.styles["Caption"]
    set_east_asia_font(caption, "宋体", 10.5, None)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def fix_structure(doc: Document) -> None:
    normalize_styles(doc)

    abstract_cn = find_paragraph(doc, lambda p: paragraph_text(p) == "摘要")
    abstract_cn.style = doc.styles["Heading 1"]

    english_title = find_paragraph(
        doc, lambda p: p.text.startswith("An Intelligent Question-Answering System for Apple Breeding Knowledge")
    )
    if not any(paragraph_text(p) == "ABSTRACT" for p in doc.paragraphs):
        abstract_en_heading = insert_paragraph_before(english_title, "ABSTRACT", "Heading 1")
        abstract_en_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    english_title.style = doc.styles["Title"]

    abstract_body = find_paragraph(doc, lambda p: p.text.startswith("Abstract:"))
    abstract_body.text = abstract_body.text.replace("Abstract:", "", 1).strip()
    abstract_body.style = doc.styles["Body Text"]

    toc_heading = find_paragraph(doc, lambda p: paragraph_text(p) in {"目 录", "目录", "目次"})
    toc_heading.text = "目次"
    toc_heading.style = doc.styles["Heading 1"]

    toc_instruction = find_paragraph(doc, lambda p: "更新目录" in p.text)
    build_toc_field(toc_instruction)

    chapter4_anchor = find_paragraph(doc, lambda p: paragraph_text(p).startswith("4.1 系统总体架构"))
    previous_non_empty = None
    for paragraph in doc.paragraphs:
        if paragraph == chapter4_anchor:
            break
        if paragraph_text(paragraph):
            previous_non_empty = paragraph
    if previous_non_empty is None or "第4章" not in previous_non_empty.text:
        insert_paragraph_before(chapter4_anchor, "第4章 系统设计与实现", "Heading 1")

    acknowledgements = find_paragraph(doc, lambda p: paragraph_text(p) == "致谢")
    acknowledgements.style = doc.styles["Heading 1"]

    biography = find_paragraph(doc, lambda p: paragraph_text(p) == "作者简历")
    biography.style = doc.styles["Heading 1"]

    for paragraph in list(doc.paragraphs):
        if paragraph_text(paragraph):
            continue
        if has_drawing(paragraph):
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name in {"font-claude-response-body", "Heading 1", "Heading 2", "p1"}:
            delete_paragraph(paragraph)


def has_drawing(paragraph: Paragraph) -> bool:
    for run in paragraph.runs:
        if run._element.xpath(".//w:drawing"):
            return True
    return False


def add_caption_after(paragraph: Paragraph, caption_text: str) -> Paragraph:
    next_element = paragraph._p.getnext()
    if next_element is not None:
        next_paragraph = Paragraph(next_element, paragraph._parent)
        if paragraph_text(next_paragraph) == caption_text:
            return next_paragraph
    caption = insert_paragraph_after(paragraph, caption_text, "Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return caption


def fix_captions(doc: Document) -> None:
    lead_in = find_paragraph(doc, lambda p: paragraph_text(p).startswith("表5-1展示了"))
    lead_in.style = doc.styles["Body Text"]
    lead_in.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table_caption_cn = find_paragraph(doc, lambda p: paragraph_text(p) == "表5-1  消融实验各配置评测结果（28题，满分10分）")
    table_caption_cn.style = doc.styles["Caption"]
    table_caption_cn.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table_caption_en = find_paragraph(doc, lambda p: paragraph_text(p).startswith("Table 5-1."))
    table_caption_en.style = doc.styles["Caption"]
    table_caption_en.alignment = WD_ALIGN_PARAGRAPH.CENTER

    body_drawings = []
    seen_summary = False
    for paragraph in doc.paragraphs:
        if paragraph_text(paragraph).startswith("摘要"):
            seen_summary = True
        if seen_summary and has_drawing(paragraph):
            body_drawings.append(paragraph)

    expected_captions = [
        "图4-1 系统总体架构图",
        "图4-2 Trait-aware 重排序评分规则示意图",
        "图4-3 Docker Compose 部署架构图",
        "图5-1 自动化评测框架",
        "图5-2 Hybrid 配置下各性状平均得分",
        "图5-3 系统典型问答示例界面（问答结果页与证据来源页）",
    ]

    if len(body_drawings) < len(expected_captions):
        raise ValueError(f"Expected at least {len(expected_captions)} body images, found {len(body_drawings)}")

    for paragraph, caption_text in zip(body_drawings[: len(expected_captions)], expected_captions):
        add_caption_after(paragraph, caption_text)


def tighten_front_matter_sections(doc: Document) -> None:
    toc_heading = find_paragraph(doc, lambda p: paragraph_text(p) == "目次")
    keywords_en = find_paragraph(doc, lambda p: paragraph_text(p).startswith("KEY WORDS:"))

    between = []
    current = keywords_en._p.getnext()
    while current is not None and current != toc_heading._p:
        between.append(Paragraph(current, keywords_en._parent))
        current = current.getnext()

    keeper = None
    for paragraph in between:
        if has_sect_pr(paragraph):
            keeper = paragraph
            break

    if keeper is None:
        keeper = insert_paragraph_after(keywords_en)
        keeper.style = doc.styles["Normal"]

    clear_paragraph(keeper)

    for paragraph in between:
        if paragraph._p is keeper._p:
            continue
        delete_paragraph(paragraph)


def split_toc_and_body(doc: Document) -> None:
    toc_heading = find_paragraph(doc, lambda p: paragraph_text(p) == "目次")
    chapter1 = find_paragraph(doc, lambda p: paragraph_text(p).startswith("第１章") or paragraph_text(p).startswith("第1章"))
    main_section_end = find_paragraph(doc, lambda p: paragraph_text(p) == "参考文献")._p.getprevious()
    # paragraph before "参考文献" still ends the main body section in the current file
    main_end_para = Paragraph(main_section_end, chapter1._parent)
    main_sect_pr = clone_sect_pr(main_end_para)
    strip_header_footer_refs(main_sect_pr)

    cursor = toc_heading._p.getnext()
    last_blank = None
    while cursor is not None and cursor != chapter1._p:
        paragraph = Paragraph(cursor, toc_heading._parent)
        if not paragraph_text(paragraph):
            last_blank = paragraph
        cursor = cursor.getnext()

    if last_blank is None:
        last_blank = insert_paragraph_before(chapter1)

    add_sect_pr(last_blank, main_sect_pr)


def split_ack_and_biography(doc: Document) -> None:
    biography = find_paragraph(doc, lambda p: paragraph_text(p) == "作者简历")
    last_sect_pr = deepcopy(doc._element.body.sectPr)
    strip_header_footer_refs(last_sect_pr)

    candidate = biography._p.getprevious()
    while candidate is not None:
        paragraph = Paragraph(candidate, biography._parent)
        if not paragraph_text(paragraph):
            add_sect_pr(paragraph, last_sect_pr)
            return
        candidate = candidate.getprevious()

    new_break = insert_paragraph_before(biography)
    add_sect_pr(new_break, last_sect_pr)


def apply_page_layout(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.6)
        section.header_distance = Cm(2)
        section.footer_distance = Cm(1.75)
        section.different_first_page_header_footer = False


def configure_headers_and_footers(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    sections = doc.sections
    if len(sections) != 10:
        raise ValueError(f"Expected 10 sections after cleanup, found {len(sections)}")

    for index, section in enumerate(sections):
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        header_paragraph = clear_story(section.header)
        footer_paragraph = clear_story(section.footer)

        if index in {0, 1, 2}:
            continue

        if index == 3:
            build_header_text(header_paragraph, "摘要")
            set_page_numbering(section, "upperRoman", 1)
        elif index == 4:
            build_header_text(header_paragraph, "ABSTRACT")
            set_page_numbering(section, "upperRoman")
        elif index == 5:
            build_header_text(header_paragraph, "目次")
            set_page_numbering(section, "upperRoman")
        elif index == 6:
            build_header_styleref(header_paragraph)
            set_page_numbering(section, "decimal", 1)
        elif index == 7:
            build_header_text(header_paragraph, "参考文献")
            set_page_numbering(section, "decimal")
        elif index == 8:
            build_header_text(header_paragraph, "致谢")
            set_page_numbering(section, "decimal")
        elif index == 9:
            build_header_text(header_paragraph, "作者简历")
            set_page_numbering(section, "decimal")

        build_page_field(footer_paragraph)


def enable_field_update_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def main() -> None:
    if not BACKUP_PATH.exists():
        BACKUP_PATH.write_bytes(THESIS_PATH.read_bytes())

    doc = Document(str(THESIS_PATH))
    fix_structure(doc)
    fix_captions(doc)
    tighten_front_matter_sections(doc)
    split_toc_and_body(doc)
    split_ack_and_biography(doc)
    apply_page_layout(doc)
    configure_headers_and_footers(doc)
    enable_field_update_on_open(doc)
    doc.save(str(THESIS_PATH))


if __name__ == "__main__":
    main()
