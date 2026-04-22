from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


CHAPTER_RE = re.compile(r"^\s*第[0-9０-９一二三四五六七八九十百]+章[\s\t　]")
SECTION_RE = re.compile(r"^\s*\d+\.\d+\s+")
SUBSECTION_RE = re.compile(r"^\s*\d+\.\d+\.\d+")
CAPTION_RE = re.compile(r"^\s*[图表]\s*\d")
LIST_LIKE_RE = re.compile(
    r"^(块大小|重叠长度|实际滑动步长|genes_|source_reported|partial|unknown|硬度/质地|采收期/成熟调控|糖度/风味|"
    r"genes_firmness|genes_color|genes_acidity|genes_harvest|genes_sugar)"
)


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    for child in settings.findall(qn("w:updateFields")):
        settings.remove(child)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def set_run_font(run, east_asia: str, western: str, size_pt: float, bold=None) -> None:
    run.font.name = western
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), western)
    r_fonts.set(qn("w:hAnsi"), western)


def set_paragraph_spacing(
    paragraph,
    *,
    align=None,
    first_indent_pt=None,
    left_indent_pt=None,
    space_before_pt=None,
    space_after_pt=None,
    line_spacing=None,
    line_spacing_rule=None,
) -> None:
    fmt = paragraph.paragraph_format
    if align is not None:
        paragraph.alignment = align
    if first_indent_pt is not None:
        fmt.first_line_indent = Pt(first_indent_pt)
    if left_indent_pt is not None:
        fmt.left_indent = Pt(left_indent_pt)
    if space_before_pt is not None:
        fmt.space_before = Pt(space_before_pt)
    if space_after_pt is not None:
        fmt.space_after = Pt(space_after_pt)
    if line_spacing is not None:
        fmt.line_spacing = line_spacing
    if line_spacing_rule is not None:
        fmt.line_spacing_rule = line_spacing_rule


def style_body(paragraph, *, english: bool = False, indent: bool = True) -> None:
    paragraph.style = "Body Text"
    set_paragraph_spacing(
        paragraph,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_indent_pt=24 if indent else 0,
        left_indent_pt=0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.5,
    )
    east = "Times New Roman" if english else "宋体"
    west = "Times New Roman"
    for run in paragraph.runs:
        set_run_font(run, east, west, 12, bold=False)


def style_keywords(paragraph, *, english: bool = False) -> None:
    paragraph.style = "Body Text"
    set_paragraph_spacing(
        paragraph,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent_pt=0,
        left_indent_pt=0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.5,
    )
    east = "Times New Roman" if english else "宋体"
    west = "Times New Roman"
    for run in paragraph.runs:
        set_run_font(run, east, west, 12, bold=False)


def style_chapter(paragraph) -> None:
    paragraph.style = "Heading 1"
    set_paragraph_spacing(
        paragraph,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent_pt=0,
        left_indent_pt=0,
        space_before_pt=12,
        space_after_pt=6,
        line_spacing=1.0,
    )
    for run in paragraph.runs:
        set_run_font(run, "黑体", "Times New Roman", 16, bold=True)


def style_section(paragraph) -> None:
    paragraph.style = "Heading 2"
    set_paragraph_spacing(
        paragraph,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent_pt=0,
        left_indent_pt=0,
        space_before_pt=12,
        space_after_pt=6,
        line_spacing=1.0,
    )
    for run in paragraph.runs:
        set_run_font(run, "黑体", "Times New Roman", 14, bold=True)


def style_subsection(paragraph) -> None:
    paragraph.style = "Heading 3"
    set_paragraph_spacing(
        paragraph,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent_pt=0,
        left_indent_pt=0,
        space_before_pt=6,
        space_after_pt=3,
        line_spacing=1.0,
    )
    for run in paragraph.runs:
        set_run_font(run, "黑体", "Times New Roman", 12, bold=True)


def style_minor_heading(paragraph) -> None:
    paragraph.style = "List Paragraph"
    set_paragraph_spacing(
        paragraph,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent_pt=0,
        left_indent_pt=0,
        space_before_pt=6,
        space_after_pt=0,
        line_spacing=1.0,
    )
    for run in paragraph.runs:
        set_run_font(run, "宋体", "Times New Roman", 12, bold=True)


def style_list_item(paragraph) -> None:
    paragraph.style = "Body Text"
    set_paragraph_spacing(
        paragraph,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent_pt=0,
        left_indent_pt=24,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.5,
    )
    for run in paragraph.runs:
        set_run_font(run, "宋体", "Times New Roman", 12, bold=False)


def style_caption(paragraph) -> None:
    paragraph.style = "Normal"
    set_paragraph_spacing(
        paragraph,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent_pt=0,
        left_indent_pt=0,
        space_before_pt=6,
        space_after_pt=6,
        line_spacing=1.0,
    )
    for run in paragraph.runs:
        set_run_font(run, "宋体", "Times New Roman", 10.5, bold=False)


def style_reference(paragraph) -> None:
    paragraph.style = "Body Text"
    set_paragraph_spacing(
        paragraph,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent_pt=-24,
        left_indent_pt=24,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=Pt(16),
        line_spacing_rule=WD_LINE_SPACING.EXACTLY,
    )
    for run in paragraph.runs:
        set_run_font(run, "宋体", "Times New Roman", 12, bold=False)


def style_preformatted(paragraph) -> None:
    paragraph.style = "Normal"
    set_paragraph_spacing(
        paragraph,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent_pt=0,
        left_indent_pt=24,
        space_before_pt=3,
        space_after_pt=3,
        line_spacing=1.15,
    )
    for run in paragraph.runs:
        set_run_font(run, "等线", "Courier New", 10.5, bold=False)


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def add_toc_before_first_chapter(doc: Document, first_chapter_index: int) -> None:
    if any("目 录" in p.text for p in doc.paragraphs):
        return
    chapter_para = doc.paragraphs[first_chapter_index]
    toc_heading = chapter_para.insert_paragraph_before("目 录")
    style_chapter(toc_heading)

    toc_field = insert_paragraph_after(toc_heading)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "打开文档后右键更新目录"
    r.append(t)
    fld.append(r)
    toc_field._p.append(fld)
    style_body(toc_field, english=False, indent=False)
    toc_field.alignment = WD_ALIGN_PARAGRAPH.CENTER

    page_break_para = insert_paragraph_after(toc_field)
    page_break_para.add_run().add_break(WD_BREAK.PAGE)


def normalize_document(doc: Document) -> None:
    paragraphs = doc.paragraphs
    first_chapter_idx = next(
        (i for i, p in enumerate(paragraphs) if CHAPTER_RE.match(p.text.strip())),
        None,
    )
    if first_chapter_idx is None:
        raise RuntimeError("未找到正文第一章，无法自动格式化。")

    add_toc_before_first_chapter(doc, first_chapter_idx)
    set_update_fields_on_open(doc)

    # Re-evaluate paragraphs after TOC insertion.
    paragraphs = doc.paragraphs
    first_chapter_idx = next(
        (i for i, p in enumerate(paragraphs) if CHAPTER_RE.match(p.text.strip())),
        None,
    )

    ref_start = next(
        (i for i, p in enumerate(paragraphs) if p.text.strip() == "参考文献"),
        None,
    )
    ack_start = next(
        (i for i, p in enumerate(paragraphs) if p.text.strip() == "致谢"),
        None,
    )

    in_cn_abstract = False
    in_en_abstract = False

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue

        if i < first_chapter_idx:
            if text.startswith("摘 要"):
                in_cn_abstract = True
                style_body(p)
            elif text.startswith("关键词"):
                in_cn_abstract = False
                style_keywords(p)
            elif text.startswith("Abstract:"):
                in_en_abstract = True
                style_body(p, english=True, indent=True)
            elif text.startswith("Keywords:"):
                in_en_abstract = False
                style_keywords(p, english=True)
            elif in_cn_abstract:
                style_body(p)
            elif in_en_abstract:
                style_body(p, english=True, indent=False)
            elif "An Intelligent Question-Answering System" in text:
                p.style = "Heading 2"
                set_paragraph_spacing(
                    p,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    first_indent_pt=0,
                    left_indent_pt=0,
                    space_before_pt=6,
                    space_after_pt=6,
                    line_spacing=1.0,
                )
                for run in p.runs:
                    set_run_font(run, "Times New Roman", "Times New Roman", 14, bold=True)
            continue

        if text in {"参考文献", "致谢"}:
            p.style = "Heading 1"
            set_paragraph_spacing(
                p,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                first_indent_pt=0,
                left_indent_pt=0,
                space_before_pt=12,
                space_after_pt=6,
                line_spacing=1.0,
            )
            for run in p.runs:
                set_run_font(run, "黑体", "Times New Roman", 16, bold=True)
            continue

        if ref_start is not None and ack_start is not None and ref_start < i < ack_start:
            style_reference(p)
            continue

        if CHAPTER_RE.match(text):
            style_chapter(p)
            continue

        if SUBSECTION_RE.match(text):
            style_subsection(p)
            continue

        if SECTION_RE.match(text):
            style_section(p)
            continue

        # Unnumbered section-level headings (e.g. "检索增强生成原理" styled as Heading 3 in source)
        if p.style.name == "Heading 3" and not SUBSECTION_RE.match(text):
            style_section(p)
            continue

        if CAPTION_RE.match(text):
            style_caption(p)
            continue

        if p.style.name == "HTML Preformatted" or text.startswith("【Level") or text.startswith("引用:") or "候选关键基因/位点" in text:
            style_preformatted(p)
            continue

        if p.style.name == "List Paragraph" and "..." not in text:
            style_minor_heading(p)
            continue

        if p.style.name == "whitespace-normal" or LIST_LIKE_RE.match(text):
            style_list_item(p)
            continue

        if p.style.name in {"font-claude-response-body", "p1", "Body Text"}:
            style_body(p)
            continue

        if p.style.name == "Normal":
            style_body(p)


def main() -> None:
    parser = argparse.ArgumentParser(description="Normalize NWAFU thesis docx formatting.")
    parser.add_argument("input_docx", type=Path)
    parser.add_argument(
        "--output",
        type=Path,
        default=None,
        help="Output path. Defaults to '<stem>_格式修订版.docx' next to the input file.",
    )
    parser.add_argument(
        "--backup",
        action="store_true",
        help="Create a '.bak.docx' backup next to the input file before writing output.",
    )
    args = parser.parse_args()

    input_path = args.input_docx.resolve()
    output_path = (
        args.output.resolve()
        if args.output
        else input_path.with_name(f"{input_path.stem}_格式修订版{input_path.suffix}")
    )

    if args.backup:
        backup_path = input_path.with_name(f"{input_path.stem}.bak{input_path.suffix}")
        shutil.copy2(input_path, backup_path)

    doc = Document(str(input_path))
    normalize_document(doc)
    doc.save(str(output_path))
    print(f"saved: {output_path}")


if __name__ == "__main__":
    main()
