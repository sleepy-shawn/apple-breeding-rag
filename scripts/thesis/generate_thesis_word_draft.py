from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path("/Users/shuaige/code/apple-breeding-rag")
TEMPLATE = ROOT / "archive/graduation-materials/templates/学姐论文模版.docx"
OUTPUT = ROOT / "葛帅-本科毕业论文初稿-v1.docx"

TITLE_CN = "面向苹果品质育种的检索增强知识问答系统设计与实现"
TITLE_EN_1 = "Design and Implementation of a Retrieval-Augmented"
TITLE_EN_2 = "Knowledge QA System for Apple Quality Breeding"

ABSTRACT_CN = [
    "摘 要：苹果果实硬度、果皮颜色、酸度、采收期和糖度等品质性状是苹果遗传育种研究中的核心目标。近年来，相关基因、QTL 和 GWAS 结果大量分散于论文正文、补充材料以及数据库表格之中，人工检索与整理效率较低，且不同研究之间证据难以快速整合，限制了苹果品质育种知识的高效利用。",
    "本研究以苹果品质育种文献、结构化基因表和 QTL/GWAS 数据为基础，设计并实现了一个面向苹果育种场景的检索增强知识问答系统。系统采用 FastAPI 与 Qdrant 构建后端服务，通过论文 PDF 解析、基因表清洗、GDR 数据转换、trait-specific collection 组织以及 rerank 策略，实现了针对硬度、颜色、酸度、采收期和糖度问题的专题化检索。为避免不同研究使用不同参考基因组造成的错误解释，系统进一步引入了坐标参考系保护机制，在保留 source-reported chr/pos 的同时，通过 reference genome 与 coordinate note 提示用户不要直接进行跨研究坐标合并。",
    "在系统实现基础上，本文构建了固定测试题集，并通过自动化评测框架对系统效果进行比较分析。结果表明，当前版本在测试集上的总体平均得分达到 8.25/10，其中 firmness、harvest 和 sugar 三类问题通过人工 curated 核心知识层得到了显著改善，系统能够稳定返回关键候选基因并给出带引用的 Level A 与 Level B 证据。研究结果说明，面向苹果品质育种场景构建 trait-specific 的 RAG 系统具有较好的可行性，可为后续苹果分子育种知识服务与智能问答研究提供工程基础。",
]

KEYWORDS_CN = "关键词：苹果育种；检索增强生成；知识问答；QTL；GWAS；品质性状"

ABSTRACT_EN = [
    "Abstract: Fruit firmness, peel color, acidity, harvest date, and sugar content are among the most important target traits in apple breeding research. In recent years, related genes, QTLs, and GWAS results have been distributed across article bodies, supplementary materials, and database tables. Manual retrieval and integration are inefficient, and evidence from different studies is difficult to compare directly.",
    "This study designed and implemented a retrieval-augmented knowledge question-answering system for apple quality breeding. The system was built on apple breeding papers, structured gene tables, and QTL/GWAS records, and implemented with FastAPI and Qdrant. Through PDF parsing, gene-table normalization, GDR conversion, trait-specific collections, and reranking strategies, the system supports targeted retrieval for firmness, color, acidity, harvest date, and sugar-related questions. To reduce interpretation risks caused by inconsistent reference genomes across studies, a coordinate guard mechanism was introduced. The system preserves source-reported chromosome and position fields while adding reference-genome notes to prevent inappropriate cross-study coordinate merging.",
    "An automated evaluation framework with a fixed benchmark set was further constructed to compare system performance. The current version achieved an overall score of 8.25/10. In particular, the firmness, harvest, and sugar categories were substantially improved after adding small curated knowledge layers, allowing the system to return key candidate genes together with traceable Level A and Level B evidence. The results indicate that a trait-specific RAG system is feasible for apple quality breeding and can provide an engineering foundation for future knowledge services and intelligent question-answering applications in horticultural breeding.",
]

KEYWORDS_EN = "Keywords: apple breeding; retrieval-augmented generation; knowledge question answering; QTL; GWAS; fruit quality traits"

TOC_LINES = [
    "第1章 文献综述.............................................................................................................................. - 1 -",
    "    1.1 苹果品质性状与分子育种研究现状....................................................................................... - 1 -",
    "    1.2 苹果品质相关 QTL/GWAS 研究进展.................................................................................... - 2 -",
    "    1.3 RAG 与农业知识问答研究进展............................................................................................ - 4 -",
    "    1.4 研究目的及意义...................................................................................................................... - 5 -",
    "第2章 研究方法与系统设计.......................................................................................................... - 6 -",
    "    2.1 数据来源与知识库构成.......................................................................................................... - 6 -",
    "    2.2 系统总体架构.......................................................................................................................... - 7 -",
    "    2.3 论文抓取与标准化处理.......................................................................................................... - 8 -",
    "    2.4 结构化基因与 GDR 数据处理.............................................................................................. - 9 -",
    "    2.5 trait-specific 检索与回答生成.............................................................................................. - 10 -",
    "    2.6 自动评测方法........................................................................................................................ - 11 -",
    "第3章 结果与分析........................................................................................................................ - 12 -",
    "    3.1 知识库构建结果...................................................................................................................... - 12 -",
    "    3.2 firmness 弱项补强结果......................................................................................................... - 13 -",
    "    3.3 QTL/GWAS 坐标参考系审计结果....................................................................................... - 14 -",
    "    3.4 自动评测结果分析.................................................................................................................. - 15 -",
    "第4章 讨论.................................................................................................................................... - 17 -",
    "第5章 结论.................................................................................................................................... - 19 -",
    "参考文献......................................................................................................................................... - 20 -",
    "致谢................................................................................................................................................. - 22 -",
]

CONTENT = {
    "第１章 文献综述": [
        ("Heading 3", "1.1 苹果品质性状与分子育种研究现状"),
        ("Body Text", "苹果果实品质是决定商品价值和消费者接受度的关键指标，其中硬度、果皮颜色、酸度、采收期和糖度是最常被关注的核心性状。这些性状通常属于典型的数量性状，既受到多基因控制，也受到环境、成熟时期和采后条件的共同影响。随着苹果基因组测序、遗传图谱构建和分子标记技术的发展，越来越多研究开始从候选基因、QTL 以及 GWAS 角度解析这些性状的遗传基础。"),
        ("Body Text", "在苹果育种实践中，研究者不仅需要关注某一候选基因是否与目标性状相关，还需要判断其证据类型是否足够直接，例如是来源于 GWAS/QTL 的显著关联，还是来自表达分析、转录调控或功能实验的间接支持。这种“证据层级”对于分子辅助育种和后续验证工作都具有重要意义。"),
        ("Heading 3", "1.2 苹果品质相关 QTL/GWAS 研究进展"),
        ("Body Text", "近年来，苹果品质性状的遗传研究积累了大量结构化信息。例如，MdNAC18、MdNAC5、MdPG、MdEXP-A1、MdERF3 和 MdERF118 等基因与果实硬度或软化过程密切相关；Ma1/MdALMT9、MdVHP1 和 MdMYB73 等基因在果实酸度调控中发挥重要作用；MdMYB1/MdMYB10、MdHY5 和 MdCOP1 等则是果皮颜色和花青素积累研究中的代表性候选基因。与此同时，不同研究也提供了丰富的 GWAS/QTL 位点、SNP 标记和候选区间信息。"),
        ("Body Text", "然而，这些信息分散于论文正文、补充表、数据库和不同格式的导出文件中，且不同研究采用的参考基因组并不完全一致。例如当前苹果相关 QTL/GWAS 数据中常见 GDDH13 v1.1、Malus domestica Whole Genome v1.0、Honeycrisp Genome v1.1.a1 以及 HFTH1 Whole Genome v1.0 等不同参考系。如果不加区分地对这些 chr/pos 进行跨研究比较，容易产生错误解释。"),
        ("Heading 3", "1.3 RAG 与农业知识问答研究进展"),
        ("Body Text", "检索增强生成（Retrieval-Augmented Generation, RAG）为专业领域知识问答提供了一种新的技术路线。其基本思想是先从知识库中检索相关内容，再结合生成模型或模板化生成模块输出答案。相比单纯依赖参数记忆的大模型，RAG 具有可追溯、可更新和可控性更强等优点。"),
        ("Body Text", "在农业和生命科学场景中，知识问答系统不仅要能够返回相关文本，更需要给出可验证的证据来源。对于苹果品质育种问题而言，理想系统应能够同时整合论文 PDF、结构化基因表、QTL/GWAS 记录以及人工整理的高价值知识片段，并根据不同性状进行专题化检索。"),
        ("Heading 3", "1.4 研究目的及意义"),
        ("Body Text", "基于上述背景，本文拟构建一个面向苹果品质育种场景的 RAG 原型系统。该系统以苹果品质文献、结构化基因表、QTL/GWAS 数据和人工 curated 核心知识层为基础，实现 trait-specific 的知识检索、证据型回答生成和自动化评测。通过这一系统，可提高苹果品质相关知识的检索效率与组织能力，为苹果分子育种研究提供可追溯的智能问答工具，也为果树育种知识服务系统的工程实现提供参考。"),
    ],
    "第２章 研究方法与系统设计": [
        ("Heading 3", "2.1 数据来源与知识库构成"),
        ("Body Text", "本研究所使用的数据主要包括两类：一类为非结构化论文数据，主要来源于苹果品质育种相关论文 PDF 及其补充材料；另一类为结构化基因/QTL/GWAS 数据，主要来源于人工整理表、GDR 转换结果以及 trait-specific curated 表。当前知识库重点覆盖硬度、颜色、酸度、采收期和糖度五类品质性状。"),
        ("Heading 3", "2.2 系统总体架构"),
        ("Body Text", "系统整体采用“数据获取与整理层—向量检索与问答服务层—前端交互与评测层”的分层架构。数据层负责论文抓取、清洗、结构化处理与 staging；服务层基于 FastAPI 和 Qdrant 构建 trait-specific 检索与回答接口；前端层提供聊天、上传和配置入口；评测层通过固定题集对各版本结果进行比较。"),
        ("Heading 3", "2.3 论文抓取与标准化处理"),
        ("Body Text", "为提高论文资料的可管理性，本研究构建了统一的工作区与 pipeline 目录规范。自动抓取结果首先进入 workspace/default/source，之后通过标准论文目录、清单文件和 ingest manifest 进行整理，再被 staging 到后端数据目录。此流程将原始抓取结果、中间状态和正式知识库区分开来，提高了后续自动化处理与重复实验的可行性。"),
        ("Heading 3", "2.4 结构化基因与 GDR 数据处理"),
        ("Body Text", "对于结构化数据，系统首先对原始 gene table 进行字段抽取与统一，随后按性状拆分为 trait-specific 子集。对于 GDR/QTL/GWAS 数据，系统构建了转换脚本和 curated layer，从 evidence_text 中尽可能抽取 candidate_gene、display_title、reference_genome 等字段。同时，为控制噪声，还引入了少量人工 curated 核心知识层，例如 firmness、harvest 和 sugar 的 golden gene layer。"),
        ("Heading 3", "2.5 trait-specific 检索与回答生成"),
        ("Body Text", "后端在接收问题后，首先检测问题所属 trait，再优先路由到对应的 trait-specific collection，例如 genes_firmness、genes_color、genes_acidity、genes_harvest 和 genes_sugar，并结合 genes_gdr_curated_* 与论文集合进行混合检索。随后通过 rerank 机制对候选结果进行重排序，提升与目标性状、关键基因和问题语义更相关的证据排名。回答部分优先输出带引用的证据型结果，并区分 Level A 直接证据与 Level B 间接证据。"),
        ("Heading 3", "2.6 自动评测方法"),
        ("Body Text", "为了对系统效果进行可重复比较，本文构建了固定评测题集。每道题设置预期基因、预期机制和证据类型，通过关键词命中、引用情况和 Level A/B 区分情况对输出结果进行打分。评测脚本会输出 results.csv、summary.md、manual_review.csv 等结果文件，便于后续版本比较和人工复核。"),
    ],
    "第３章 结果与分析": [
        ("Heading 3", "3.1 知识库构建结果"),
        ("Body Text", "当前系统已建立论文集合、通用基因集合、trait-specific gene collections 以及 GDR curated collections。知识库覆盖硬度、颜色、酸度、采收期和糖度等核心苹果品质性状，并形成了可用于问答和评测的统一数据组织框架。"),
        ("Heading 3", "3.2 firmness 弱项补强结果"),
        ("Body Text", "在系统迭代过程中，firmness 方向一度是评测中的主要弱项，特别是与 Honeycrisp 脆度和质地相关的问题，原始检索结果容易被大量泛化的 texture/QTL 记录占据前排，无法稳定返回 MdNAC18、MdPG 或 MdEXP-A1 等更符合科研语境的核心候选基因。为解决该问题，本文构建了 genes_firmness_curated.csv，将 MdNAC18、MdNAC5、MdPG、MdEXP-A1、MdEXP、MdERF3、MdERF118、MdACS1、MdACO1 和 MdPAE10 等高价值证据整理为小规模人工核心知识层。"),
        ("Body Text", "在引入 firmness curated layer 后，Honeycrisp/texture 相关问题的回答显著改善，系统能够稳定命中 MdNAC18、MdEXP-A1 和 MdPG 等关键基因，避免被泛化位点完全抢占前排。"),
        ("Heading 3", "3.3 QTL/GWAS 坐标参考系审计结果"),
        ("Body Text", "针对 QTL/GWAS 数据中不同研究参考基因组不一致的问题，本文新增了参考系审计脚本，对 gene/GDR CSV 中的 chr/pos 与 reference_genome 字段进行统计分析。结果表明，现有数据中包含 GDDH13 v1.1、Malus domestica Whole Genome v1.0、Honeycrisp Genome v1.1.a1、HFTH1 Whole Genome v1.0 以及少量 Pyrus pyrifolia 参考系。基于此，系统采用“最小安全策略”，仅保留 source-reported 坐标与参考基因组提示，不进行 liftover，也不进行跨研究坐标级共定位合并。"),
        ("Heading 3", "3.4 自动评测结果分析"),
        ("Body Text", "在当前最佳版本 baseline_firmness_texture_curated 中，系统总体平均得分达到 8.25/10。其中 firmness 为 8.6/10，color 为 8.0/10，acidity 为 8.2/10，harvest 为 8.5/10，sugar 为 9.0/10。结果说明，trait-specific collection 与人工 curated layer 的结合显著提升了苹果品质问题的检索稳定性和证据可追溯性。"),
        ("Body Text", "与早期版本相比，当前版本在 firmness 方向提升最为明显，说明针对弱项构建小规模高质量知识层，比盲目扩大原始数据规模更有效。与此同时，坐标参考系保护机制并未降低当前 baseline 得分，表明在不做复杂 liftover 的前提下，仍可以通过元数据提示保证结论的安全性。"),
    ],
    "第４章 讨论": [
        ("Body Text", "本文构建的苹果育种 RAG 原型系统，已经能够较稳定地回答苹果品质育种中的 trait-specific 问题，并以可追溯的形式输出证据。其工程价值主要体现在三个方面：第一，系统实现了论文 PDF、结构化 gene table 和 GDR/QTL/GWAS 数据的统一接入；第二，通过 trait-specific collection 和 rerank 机制，降低了跨性状噪声对结果的影响；第三，通过人工 curated layer 对关键弱项进行针对性补强，提高了科研问答场景中的实用性。"),
        ("Body Text", "但需要指出的是，当前系统仍然不能被视为完全自动的知识发现系统。部分性能提升来自人工 curated/golden gene layer，因此论文中应将其表述为“人工确认核心知识层”或“curated knowledge layer”，而非系统自动挖掘出的新结论。此外，自动评分虽然能够有效支持版本比较，但仍不能替代导师或领域专家的人工判断。"),
        ("Body Text", "从后续研究角度看，最有价值的工作并不是继续盲目增加文献数量，而是进一步补齐 curated 数据的 DOI、PMID、supplement table 和 evidence strength 等来源字段，建立更加严谨的老师确认闭环。同时，在具备充分 marker 映射信息和 reference build 信息之后，才能进一步探索真正可靠的坐标级 liftover 与共定位分析。"),
    ],
    "第５章 结论": [
        ("Body Text", "本文围绕苹果品质育种知识问答需求，设计并实现了一个面向苹果育种场景的检索增强知识问答系统。系统集成论文 PDF、结构化 gene table、QTL/GWAS 数据和人工 curated 核心知识层，通过 FastAPI、Qdrant、trait-specific collection、rerank 和自动化评测框架，实现了对硬度、颜色、酸度、采收期和糖度等问题的知识检索与回答。"),
        ("Body Text", "实验结果表明，当前系统在固定测试题集上取得了较稳定的性能，整体得分达到 8.25/10，其中 firmness、harvest 和 sugar 等方向得到明显改善。系统不仅能够返回相关候选基因，还能够给出带引用的 Level A 与 Level B 证据，并对 QTL/GWAS 坐标参考系不一致问题进行保护性提示。"),
        ("Body Text", "总体而言，本研究完成了一个具备实际展示价值和继续扩展潜力的苹果育种 RAG 原型系统，为后续苹果分子育种知识服务、智能问答和领域知识工程研究提供了可行的技术基础。"),
    ],
}

REFERENCES = [
    "Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. Advances in Neural Information Processing Systems, 2020.",
    "Guu K, Lee K, Tung Z, et al. REALM: Retrieval-Augmented Language Model Pre-Training[C]. International Conference on Machine Learning, 2020.",
    "Watts S, Migicovsky Z, Myles S. Large-scale apple GWAS reveals NAC18.1 as a master regulator of ripening traits[J]. Fruit Research, 2023, 3:32.",
    "Longhi S, Hamblin M T, Trainotti L, et al. A candidate gene based approach validates Md-PG1 as the main responsible for a QTL impacting fruit texture in apple (Malus domestica Borkh.)[J]. BMC Plant Biology, 2013, 13:37.",
    "Chagné D, Dayatilake D, Diack R, et al. Genetic and environmental control of fruit maturation, dry matter and firmness in apple (Malus domestica Borkh.)[J]. Horticulture Research, 2014, 1:14046.",
    "Jung M, Keller B, Roth M, et al. Genetic architecture and genomic predictive ability of apple quantitative traits across environments[J]. 2022.",
    "Liu Y, et al. Natural variation in MdNAC5 contributes to fruit firmness and ripening divergence in apple[J]. 2024.",
    "Allelic variation in an expansin, MdEXP-A1, contributes to flesh firmness at harvest in apples[J]. 2024.",
    "Role of MdERF3 and MdERF118 natural variations in apple flesh firmness crispness retainability and development of QTL-based genomics-assisted prediction[J]. 2024.",
    "Natural variations in MdPAE10 contribute to prolonged apple fruit shelf life[J]. 2024.",
    "Minamikawa M F, Kunihisa M, Moriya S, et al. Genomic prediction and genome-wide association study using combined genotypic data from different genotyping systems: application to apple fruit quality traits[J]. Horticulture Research, 2024.",
    "Moriya S, Kunihisa M, Okada K, et al. Identification of QTLs for Flesh Mealiness in Apple (Malus domestica Borkh.)[J]. The Horticulture Journal, 2017, 86(2):159-170.",
]

ACK = (
    "感谢指导教师在选题、系统设计和论文写作过程中给予的指导与帮助。"
    "感谢课题组同学在苹果文献整理、基因数据收集和系统测试中的支持。"
    "同时感谢家人与朋友在毕业设计期间给予的理解与鼓励。"
)


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def delete_table(table) -> None:
    tbl = table._element
    parent = tbl.getparent()
    if parent is not None:
        parent.remove(tbl)


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text


STYLE_MAP = {
    "Heading 2": "Heading2",
    "Heading 3": "Heading3",
    "Body Text": "BodyText",
    "List Paragraph": "ListParagraph",
}


def add_styled_paragraph(doc: Document, text: str, style_name: str):
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles[STYLE_MAP.get(style_name, style_name)]
    paragraph.text = text
    return paragraph


def add_body(doc: Document) -> None:
    first = True
    for chapter, blocks in CONTENT.items():
        if not first:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        first = False
        add_styled_paragraph(doc, chapter, "Heading 2")
        for style, text in blocks:
            add_styled_paragraph(doc, text, style)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_styled_paragraph(doc, "参考文献", "Heading 2")
    for idx, ref in enumerate(REFERENCES, start=1):
        add_styled_paragraph(doc, f"[{idx}] {ref}", "Body Text")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_styled_paragraph(doc, "致谢", "Heading 2")
    add_styled_paragraph(doc, ACK, "Body Text")


def main() -> None:
    doc = Document(str(TEMPLATE))

    # Cover page
    set_paragraph_text(doc.paragraphs[0], "学号：2022010637")
    set_paragraph_text(doc.paragraphs[6], " 2026 届本科生毕业论文")
    set_paragraph_text(doc.paragraphs[11], TITLE_CN)
    set_paragraph_text(doc.paragraphs[23], "协助指导教师：")
    set_paragraph_text(doc.paragraphs[25], "完 成 日 期：\t2026年05月")

    cover_table = doc.tables[0]
    cover_table.cell(0, 4).text = "园艺学院"
    cover_table.cell(1, 4).text = "园艺专业"
    cover_table.cell(2, 4).text = "2022级05班"
    cover_table.cell(3, 4).text = "葛帅"
    cover_table.cell(4, 4).text = "待补充"

    # Abstracts
    set_paragraph_text(doc.paragraphs[46], TITLE_CN)
    set_paragraph_text(doc.paragraphs[48], ABSTRACT_CN[0])
    set_paragraph_text(doc.paragraphs[49], ABSTRACT_CN[1])
    set_paragraph_text(doc.paragraphs[50], ABSTRACT_CN[2])
    set_paragraph_text(doc.paragraphs[52], KEYWORDS_CN)

    set_paragraph_text(doc.paragraphs[56], TITLE_EN_1)
    set_paragraph_text(doc.paragraphs[58], TITLE_EN_2)
    set_paragraph_text(doc.paragraphs[59], ABSTRACT_EN[0])
    set_paragraph_text(doc.paragraphs[60], ABSTRACT_EN[1])
    set_paragraph_text(doc.paragraphs[61], ABSTRACT_EN[2])
    set_paragraph_text(doc.paragraphs[63], KEYWORDS_EN)

    # Table of contents placeholder
    for idx, line in zip(range(69, 69 + len(TOC_LINES)), TOC_LINES):
        set_paragraph_text(doc.paragraphs[idx], line)
    for idx in range(69 + len(TOC_LINES), 101):
        set_paragraph_text(doc.paragraphs[idx], "")

    # Remove old body content
    for paragraph in reversed(doc.paragraphs[103:]):
        delete_paragraph(paragraph)
    for table in list(doc.tables)[2:]:
        delete_table(table)

    add_body(doc)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
