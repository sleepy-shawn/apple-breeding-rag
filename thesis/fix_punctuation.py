"""
fix_punctuation.py
------------------
将论文 docx 中混用的英文标点符号批量替换为中文标点。
仅在中文上下文中替换，保留英文引用（如 Lewis et al., 2020）中的英文标点。

策略：在段落全文上做上下文判断，再把修正后的字符分配回各 Run，保留格式。
"""

import re
from pathlib import Path
from docx import Document

SRC = Path(__file__).parent / "葛帅毕业论文.docx"
DST = Path(__file__).parent / "葛帅毕业论文_标点修正.docx"


# ── 辅助函数 ──────────────────────────────────────────────

def is_chinese(ch: str) -> bool:
    """判断是否是中文字符（CJK统一汉字区间）。"""
    return "一" <= ch <= "鿿" or "　" <= ch <= "〿"


def has_chinese(text: str) -> bool:
    return any(is_chinese(c) for c in text)


def is_reference_para(text: str) -> bool:
    """
    粗判：英文参考文献行（以大写字母+小写字母+空格+大写字母开头）。
    此类行不做标点替换。
    """
    stripped = text.strip()
    return bool(re.match(r"^[A-Z][a-z]+ [A-Z]", stripped))


# ── 核心替换：在整段文本上做上下文感知替换 ────────────────

def smart_replace_text(text: str) -> str:
    """
    对整段文本做中文上下文敏感的标点替换。
    保留：英文字母旁的逗号（引用格式）、URL 等。
    """
    if not has_chinese(text):
        return text

    chars = list(text)
    n = len(chars)

    def prev_ch(i):
        return chars[i - 1] if i > 0 else ""

    def next_ch(i):
        return chars[i + 1] if i < n - 1 else ""

    for i in range(n):
        ch = chars[i]
        p, nx = prev_ch(i), next_ch(i)

        if ch == ",":
            # 替换条件：前或后紧邻中文字符
            if is_chinese(p) or is_chinese(nx):
                chars[i] = "，"

        elif ch == ";":
            if is_chinese(p) or is_chinese(nx):
                chars[i] = "；"

        elif ch == "(":
            # 前一字符是中文，或后一字符是中文
            if is_chinese(p) or is_chinese(nx):
                chars[i] = "（"

        elif ch == ")":
            # 后一字符是中文，或前一字符是中文
            if is_chinese(nx) or is_chinese(p):
                chars[i] = "）"

    return "".join(chars)


def fix_extra_corrections(text: str) -> str:
    """额外的针对性修正。"""
    # 物种名空格：Malus×domestica → Malus × domestica
    text = re.sub(r"Malus×domestica", "Malus × domestica", text)
    # 图/表序号空格：图 5- 3 → 图5-3
    text = re.sub(r"(?<=[图表])\s*(\d+)\s*-\s*(\d+)", r"\1-\2", text)
    # 补全 sentence-transformers 引用（如果还未添加）
    old = "sentence-transformers 系列的轻量化嵌入模型完成向量化，生成384维"
    new = "sentence-transformers 系列的轻量化嵌入模型完成向量化（Sentence Transformers, 2026），生成384维"
    text = text.replace(old, new)
    return text


# ── 段落级处理：保留 Run 格式 ────────────────────────────

def process_paragraph_runs(para) -> bool:
    """
    1. 收集所有 run 的文本，拼成完整段落文本
    2. 对完整文本做替换
    3. 把替换后的字符按原 run 长度重新分配回各 run
    返回 True 表示有改动。
    """
    runs = para.runs
    if not runs:
        return False

    # 拼完整文本
    full_original = "".join(r.text for r in runs)
    if not full_original.strip():
        return False

    # 跳过参考文献
    if is_reference_para(full_original):
        return False

    # 做替换
    full_new = smart_replace_text(full_original)
    full_new = fix_extra_corrections(full_new)

    if full_new == full_original:
        return False

    # 按原 run 长度把修改后的字符分配回各 run
    pos = 0
    for run in runs:
        length = len(run.text)
        run.text = full_new[pos: pos + length]
        pos += length

    return True


# ── 主程序 ────────────────────────────────────────────────

def main():
    print(f"读取原始文件：{SRC.name}")
    doc = Document(str(SRC))

    modified_paras = 0

    # 处理正文段落
    for para in doc.paragraphs:
        if process_paragraph_runs(para):
            modified_paras += 1

    # 处理表格单元格内段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if process_paragraph_runs(para):
                        modified_paras += 1

    doc.save(str(DST))
    print(f"✅ 完成！共修改了 {modified_paras} 个段落/单元格")
    print(f"输出文件：{DST.name}")

    # ── 验证：抽查几段 ──
    print("\n--- 抽查验证 ---")
    doc2 = Document(str(DST))
    shown = 0
    for para in doc2.paragraphs:
        if shown >= 6:
            break
        txt = para.text
        # 找含中文且含中文逗号的段落，说明替换已生效
        if has_chinese(txt) and "，" in txt and len(txt) > 30:
            # 检查是否还残留英文逗号（中文旁）
            residual = sum(
                1 for i, c in enumerate(txt)
                if c == ","
                and (is_chinese(txt[i - 1]) if i > 0 else False
                     or is_chinese(txt[i + 1]) if i < len(txt) - 1 else False)
            )
            marker = "⚠️ 仍有残留" if residual else "✓"
            print(f"  {marker} {txt[:90]}...")
            shown += 1

    # 统计残留英文逗号数（中文上下文中）
    total_residual = 0
    doc3 = Document(str(DST))
    for para in doc3.paragraphs:
        txt = para.text
        if not has_chinese(txt):
            continue
        for i, c in enumerate(txt):
            if c == ",":
                p = txt[i - 1] if i > 0 else ""
                nx = txt[i + 1] if i < len(txt) - 1 else ""
                if is_chinese(p) or is_chinese(nx):
                    total_residual += 1
    print(f"\n残留英文逗号（中文上下文中）：{total_residual} 处")


if __name__ == "__main__":
    main()
