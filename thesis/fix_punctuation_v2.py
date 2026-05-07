"""
fix_punctuation_v2.py
=====================
论文 v2 补丁脚本——基于 awesome-ai-research-writing 的「表达润色」方法论。

输入：葛帅毕业论文_标点修正.docx
输出：葛帅毕业论文_v2.docx

修复策略遵循"尊重原著、克制修改"原则：
  · A 类硬错误：截断、括号不匹配、缺标点、英文标点 → 全部修复
  · B 类风格改进：仅改动确有问题的（口语化、语法卡顿、参考文献格式）
  · 不为追求形式变化做强行同义替换
"""

import re
from pathlib import Path
from docx import Document

SRC = Path(__file__).parent / "葛帅毕业论文_标点修正.docx"
DST = Path(__file__).parent / "葛帅毕业论文_v2.docx"

# ── 开关（B 类风格改动可独立关闭）─────────────────────────
APPLY_B1 = True   # "张冠李戴" → "指代错误"
APPLY_B2 = True   # "工具有三个核心要求" → "工具提出三个核心要求"
APPLY_B3 = False  # 第 2.2.6 节并列项重写（默认关）
APPLY_B5 = True   # 中文参考文献作者间分隔符统一


# ── 辅助 ────────────────────────────────────────────────
def is_chinese(ch: str) -> bool:
    return "一" <= ch <= "鿿" or "　" <= ch <= "〿"


def has_chinese(text: str) -> bool:
    return any(is_chinese(c) for c in text)


def is_reference_para(text: str) -> bool:
    """英文参考文献（以 'Author A, Author B' 格式开头）。"""
    stripped = text.strip()
    return bool(re.match(r"^[A-Z][a-z]+ [A-Z]", stripped))


def is_chinese_reference_para(text: str) -> bool:
    """中文参考文献（中文作者名 + 年份 + 期刊标记[J]/[C]/[M]/[EB/OL]）。"""
    stripped = text.strip()
    if not has_chinese(stripped):
        return False
    # 包含年份点和标记
    has_year_marker = bool(re.search(r"\d{4}\.\s.+\[(J|C|M|D|EB/OL|PP/OL)\]", stripped))
    # 以中文人名起头（首字符为中文）
    starts_with_chinese = is_chinese(stripped[0]) if stripped else False
    return starts_with_chinese and has_year_marker


# ── A 类硬错误：精确字符串替换 ──────────────────────────
A_FIXES = [
    # A1: 摘要末尾截断
    ("成熟证据链相吻", "成熟证据链相吻合。"),

    # A2: 第3.2.3节后半段被删
    (
        "适合在本地部署环境下处理苹果育种相关的英文文献片段。向量连同对应",
        "适合在本地部署环境下处理苹果育种相关的英文文献片段。向量连同对应的元数据一并写入向量数据库的论文检索索引，供后续检索调用。"
    ),

    # A3: 第5.3.4节多余左括号 + 缺句号
    (
        "是以完全失去引用追溯能力和证据分层能力（为代价的，A0 模式下回答虽然流畅",
        "是以完全失去引用追溯能力和证据分层能力为代价的。A0 模式下回答虽然流畅"
    ),

    # A5: 第5.3.4节缺句号
    (
        "正是为了避免这种单一维度评估的偏误只有在四个维度都达到合理水平时",
        "正是为了避免这种单一维度评估的偏误。只有在四个维度都达到合理水平时"
    ),

    # A6: 第5.4.1节缺冒号
    (
        "之间存在乙烯生物合成通路上的层级关系MdNAC18.1 作为上游转录因子",
        "之间存在乙烯生物合成通路上的层级关系：MdNAC18.1 作为上游转录因子"
    ),

    # A7: 第5.4引文括号不匹配
    (
        "（Migicovsky et al., 2021)",
        "（Migicovsky et al., 2021）"
    ),

    # A8: 英文摘要多余空格
    (
        "the best performance , consistent with",
        "the best performance, consistent with"
    ),

    # A12: 微卫星 → 串联重复（事实修正）
    (
        "MdMYB10 启动子区的 R6 微卫星重复",
        "MdMYB10 启动子区的 R6 串联重复"
    ),

    # A13: 第5.3.3节两句之间缺标点（引号是 ASCII 双引号 U+0022）
    (
        '结构化基因记录擅长提供"是什么"论文上下文则擅长提供',
        '结构化基因记录擅长提供"是什么"，论文上下文则擅长提供'
    ),
]

# B1: 张冠李戴 → 指代错误
B1_FIXES = [
    ("纯大语言模型容易出现张冠李戴的事实错误", "纯大语言模型容易出现指代错误的事实偏差"),
]

# B2: 工具有 → 工具提出
B2_FIXES = [
    (
        "苹果育种家在实际工作中对知识服务工具有三个核心要求",
        "苹果育种家在实际工作中对知识服务工具提出三个核心要求"
    ),
]

# B3: 第2.2.6节并列项不平行（默认关闭）
B3_FIXES = [
    (
        "这五类性状共同覆盖了苹果商品价值的主要维度，口感（硬度、糖度、酸度）、外观、采后表现与市场窗口，具备产业代表性。",
        "这五类性状共同覆盖了苹果商品价值的主要维度，涵盖口感（硬度、糖度、酸度）、外观、采后表现与市场窗口四个层面，具备产业代表性。"
    ),
]


# ── 中文上下文标点替换（A9, A10）────────────────────────
def smart_replace_punctuation(text: str) -> str:
    """对 :  ?  做中文上下文敏感替换。"""
    if not has_chinese(text):
        return text
    chars = list(text)
    n = len(chars)
    for i, ch in enumerate(chars):
        prev = chars[i - 1] if i > 0 else ""
        nxt = chars[i + 1] if i < n - 1 else ""

        if ch == ":":
            # 排除：URL（前面有 http、https、www、ftp）、坐标格式（如 Chr03:）
            # 简单判断：前 6 字符内有 http/www，则跳过
            window = "".join(chars[max(0, i - 6):i])
            if "http" in window or "ftp" in window or "www" in window:
                continue
            # 坐标格式 Chr\d+: 保留英文冒号（科学惯例）
            if i >= 4 and re.match(r"[Cc]hr\d+", "".join(chars[max(0, i - 5):i])):
                continue
            if is_chinese(prev) or is_chinese(nxt):
                chars[i] = "："

        elif ch == "?":
            if is_chinese(prev) or is_chinese(nxt):
                chars[i] = "？"

    return "".join(chars)


# ── A4: 中左+英右括号修复 ────────────────────────────────
def fix_mismatched_paren(text: str) -> str:
    """
    （...) → （...）
    匹配以中文左括号开头、内部不含括号、以英文右括号结尾的字符串。
    """
    return re.sub(r"（([^（）()]*?)\)", r"（\1）", text)


# ── A10: 预置问题列表 ───────────────────────────────────
def fix_preset_questions(text: str) -> str:
    """
    "...?" ，X性状的Y → "...？"——X性状的Y
    把英文问号换成中文，引号后多余空格 + 逗号改成"——"破折号。
    """
    # 匹配模式："xxx?" ，
    text = re.sub(
        r'(["“][^"”]*?)\?(["”])\s*，',
        lambda m: f"{m.group(1)}？{m.group(2)}——",
        text,
    )
    return text


# ── B5: 中文参考文献作者间分隔符 ────────────────────────
def fix_chinese_reference(text: str) -> str:
    """
    中文参考文献中将 '， ' 作者间分隔符改为 ', '（半角逗号+空格）。
    仅改动到第一个数字（年份）出现之前的部分。
    """
    if not is_chinese_reference_para(text):
        return text
    # 只改前半段（年份之前）
    m = re.search(r"(\d{4})\.\s", text)
    if not m:
        return text
    head = text[:m.start()]
    tail = text[m.start():]
    # 将"， "改为", "
    head_fixed = head.replace("， ", ", ")
    return head_fixed + tail


# ── 段落处理（保留 Run 格式）─────────────────────────────
def apply_all_fixes(text: str) -> str:
    """对一段文本应用所有 A、B 类修复。"""
    new = text

    # A 类硬错误：精确字符串替换
    for old, new_str in A_FIXES:
        new = new.replace(old, new_str)

    # A4: 中左英右括号
    new = fix_mismatched_paren(new)

    # A9: 英文冒号 + A10: 英文问号（中文上下文）
    new = smart_replace_punctuation(new)

    # A10: 预置问题特殊格式
    new = fix_preset_questions(new)

    # B 类（按开关）
    if APPLY_B1:
        for old, new_str in B1_FIXES:
            new = new.replace(old, new_str)
    if APPLY_B2:
        for old, new_str in B2_FIXES:
            new = new.replace(old, new_str)
    if APPLY_B3:
        for old, new_str in B3_FIXES:
            new = new.replace(old, new_str)
    if APPLY_B5:
        new = fix_chinese_reference(new)

    return new


def process_paragraph_runs(para) -> tuple[bool, str]:
    """
    返回 (是否修改, 修改后的文本)。
    保持 run 结构，将修改后的字符按原 run 长度回填。
    """
    runs = para.runs
    if not runs:
        return False, para.text

    full_old = "".join(r.text for r in runs)
    if not full_old.strip():
        return False, full_old

    # 跳过英文参考文献（只受 A 类英文修复影响）
    if is_reference_para(full_old):
        # 仍要做 A8 这类英文修复
        full_new = full_old
        for old, new_str in A_FIXES:
            full_new = full_new.replace(old, new_str)
        if full_new == full_old:
            return False, full_old
    else:
        full_new = apply_all_fixes(full_old)
        if full_new == full_old:
            return False, full_old

    # 长度可能变化，需重新分配到 runs
    if len(full_new) == len(full_old):
        # 等长：直接按原长度切片回填
        pos = 0
        for run in runs:
            length = len(run.text)
            run.text = full_new[pos: pos + length]
            pos += length
    else:
        # 长度变化：把所有新文本放进第一个 run，其余清空
        runs[0].text = full_new
        for r in runs[1:]:
            r.text = ""

    return True, full_new


# ── 主程序 ────────────────────────────────────────────────
def main():
    print(f"读取：{SRC.name}")
    print(f"风格开关：B1={APPLY_B1}  B2={APPLY_B2}  B3={APPLY_B3}  B5={APPLY_B5}")
    print()

    doc = Document(str(SRC))
    modified = 0
    samples = []

    for i, para in enumerate(doc.paragraphs):
        old = para.text
        changed, new = process_paragraph_runs(para)
        if changed:
            modified += 1
            if len(samples) < 8:
                samples.append((i, old[:60], new[:60]))

    # 表格内段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    changed, _ = process_paragraph_runs(para)
                    if changed:
                        modified += 1

    doc.save(str(DST))
    print(f"✅ 共修改 {modified} 个段落/单元格")
    print(f"输出：{DST.name}")

    # ── 验证关键修复 ──────────────────────────────────────
    print()
    print("=" * 60)
    print("关键修复验证：")
    print("=" * 60)
    doc2 = Document(str(DST))
    checks = [
        ("A1 摘要末尾", "成熟证据链相吻合。"),
        ("A2 3.2.3节末尾", "供后续检索调用。"),
        ("A3 5.3.4节括号", "证据分层能力为代价的。"),
        ("A5 5.3.4节句号", "评估的偏误。只有在四个"),
        ("A6 5.4.1节冒号", "层级关系：MdNAC18.1"),
        ("A7 引文括号", "（Migicovsky et al., 2021）"),
        ("A8 英文摘要", "performance, consistent"),
        ("A12 串联重复", "R6 串联重复"),
        ("A13 5.3.3两句缺标点", '"是什么"，论文上下文则擅长'),
    ]
    if APPLY_B1:
        checks.append(("B1 张冠李戴→指代错误", "指代错误的事实偏差"))
    if APPLY_B2:
        checks.append(("B2 工具有→工具提出", "工具提出三个核心要求"))

    full_doc_text = "\n".join(p.text for p in doc2.paragraphs)
    for name, marker in checks:
        ok = marker in full_doc_text
        print(f"  {'✓' if ok else '✗'} {name}: {marker[:40]!r}")

    # 残留英文标点统计
    print()
    print("残留检测（中文上下文中的英文标点）：")
    residual = {":": 0, "?": 0, ",": 0}
    for p in doc2.paragraphs:
        t = p.text
        if not has_chinese(t):
            continue
        if is_reference_para(t):
            continue
        for i, c in enumerate(t):
            if c in residual:
                prev = t[i - 1] if i > 0 else ""
                nxt = t[i + 1] if i < len(t) - 1 else ""
                # 排除 URL 和坐标
                if c == ":":
                    window = t[max(0, i - 6):i]
                    if "http" in window or "ftp" in window or "www" in window:
                        continue
                    if i >= 4 and re.match(r"[Cc]hr\d+", t[max(0, i - 5):i]):
                        continue
                if is_chinese(prev) or is_chinese(nxt):
                    residual[c] += 1
    for c, n in residual.items():
        print(f"  '{c}' 残留：{n} 处")


if __name__ == "__main__":
    main()
