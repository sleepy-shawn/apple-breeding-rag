"""
fix_citations_v3.py
===================
论文 v3 引用修复脚本。

输入：葛帅毕业论文_v2.docx
输出：葛帅毕业论文_v3.docx

修复项：
  C1：5.1.2 节 DeepSeek Chat API 后补 `（DeepSeek, 2026）` 引用
  C2：4 处英左+中右括号统一为中文括号
  C3：Espley 2009 引用删除多余的期刊名/卷期/页码
  C4：3 处多重引用英文分号 `;` → 中文 `；`
  C5（可选）：Migicovsky 等"等人"前补空格
"""

import re
from pathlib import Path
from docx import Document

SRC = Path(__file__).parent / "葛帅毕业论文_v2.docx"
DST = Path(__file__).parent / "葛帅毕业论文_v3.docx"

# ── 开关 ──────────────────────────────────────────────
APPLY_C5 = True   # Migicovsky 等"等人"前补空格


# ── 辅助 ──────────────────────────────────────────────
def is_chinese(c):
    return "一" <= c <= "鿿" or "　" <= c <= "〿"


def has_chinese(t):
    return any(is_chinese(c) for c in t)


def is_reference_para(t):
    s = t.strip()
    return bool(re.match(r"^[A-Z][a-z]+ [A-Z]", s))


# ── C1：DeepSeek 引用补全 ──────────────────────────────
# 仅在 5.1.2 节首次出现处补引用，第二次提到（5.5.2）通常不再重复引用
C1_FIXES = [
    (
        "本研究采用LLM-as-Judge 方法，由 DeepSeek Chat API 依据每道题的 expected_mechanism 字段对系统回答进行自动打分",
        "本研究采用LLM-as-Judge 方法，由 DeepSeek Chat API（DeepSeek, 2026）依据每道题的 expected_mechanism 字段对系统回答进行自动打分"
    ),
]

# ── C2：4 处英左+中右括号 ─────────────────────────────
C2_FIXES = [
    ("v1.0(Velasco et al., 2010）",   "v1.0（Velasco et al., 2010）"),
    ("v1.1(Daccord et al., 2017）",   "v1.1（Daccord et al., 2017）"),
    ("v1.0(Zhang et al., 2019）",     "v1.0（Zhang et al., 2019）"),
    ("v1.1(Khan et al., 2022）",      "v1.1（Khan et al., 2022）"),
]

# ── C3：Espley 2009 引用删除冗余期刊信息 ──────────────
C3_FIXES = [
    (
        "（Espley et al., 2009, Plant Cell 21: 168–183）",
        "（Espley et al., 2009）"
    ),
]

# ── C4：多重引用英文分号→中文分号（仅在引用括号内）─
def fix_multi_citation_semicolon(text: str) -> str:
    """
    匹配 (Author et al., YYYY; Author et al., YYYY) 这种引用括号
    把内部的英文 `;` 改为中文 `；`
    """
    def replace_in_paren(m):
        inner = m.group(1)
        # 仅当括号内符合"作者+年份"格式时才替换
        if re.search(r'(?:[A-Za-z]+\s*(?:et\s*al\.|and\s+[A-Z][a-zA-Z]+)?[,，]\s*\d{4})', inner):
            inner = inner.replace("; ", "；").replace(";", "；")
        return f"（{inner}）"

    # 处理全角括号引用
    text = re.sub(r"（([^（）()]{5,200})）", replace_in_paren, text)
    return text


# ── C5：Migicovsky 等人 → Migicovsky 等人 (确保有空格) ────
def fix_migicovsky_spacing(text: str) -> str:
    if not APPLY_C5:
        return text
    # "Migicovsky等人" → "Migicovsky 等人"（仅当前面是英文字母，后面是"等"）
    text = re.sub(r"([A-Za-z])(等人?)（", r"\1 \2（", text)
    return text


# ── 主处理 ────────────────────────────────────────────
def apply_all_fixes(text: str) -> str:
    new = text
    for old, replacement in C1_FIXES:
        new = new.replace(old, replacement)
    for old, replacement in C2_FIXES:
        new = new.replace(old, replacement)
    for old, replacement in C3_FIXES:
        new = new.replace(old, replacement)
    new = fix_multi_citation_semicolon(new)
    new = fix_migicovsky_spacing(new)
    return new


def process_paragraph_runs(para):
    runs = para.runs
    if not runs:
        return False
    full_old = "".join(r.text for r in runs)
    if not full_old.strip():
        return False
    if is_reference_para(full_old):
        # 参考文献本身不动
        return False

    full_new = apply_all_fixes(full_old)
    if full_new == full_old:
        return False

    if len(full_new) == len(full_old):
        pos = 0
        for run in runs:
            length = len(run.text)
            run.text = full_new[pos:pos + length]
            pos += length
    else:
        runs[0].text = full_new
        for r in runs[1:]:
            r.text = ""
    return True


def main():
    print(f"读取：{SRC.name}")
    print(f"开关：C5={APPLY_C5}")
    print()

    doc = Document(str(SRC))
    modified = 0

    for para in doc.paragraphs:
        if process_paragraph_runs(para):
            modified += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if process_paragraph_runs(para):
                        modified += 1

    doc.save(str(DST))
    print(f"✅ 共修改 {modified} 个段落")
    print(f"输出：{DST.name}\n")

    # ── 验证 ────────────────────────────────────────
    print("=" * 60)
    print("修复验证：")
    print("=" * 60)
    doc2 = Document(str(DST))
    full = "\n".join(p.text for p in doc2.paragraphs)

    checks = [
        ("C1 DeepSeek 引用补全",
         "DeepSeek Chat API（DeepSeek, 2026）"),
        ("C2.1 Velasco 括号统一",
         "v1.0（Velasco et al., 2010）"),
        ("C2.2 Daccord 括号统一",
         "v1.1（Daccord et al., 2017）"),
        ("C2.3 Zhang 括号统一",
         "v1.0（Zhang et al., 2019）"),
        ("C2.4 Khan 括号统一",
         "v1.1（Khan et al., 2022）"),
        ("C3 Espley 2009 简化",
         "（Espley et al., 2009）；在酸度性状"),
        ("C4 多重引用中文分号",
         "（Karpukhin et al., 2020；Khattab and Zaharia, 2020）"),
    ]
    for name, marker in checks:
        ok = marker in full
        print(f"  {'✓' if ok else '✗'} {name}")
        if not ok:
            print(f"      期望: {marker[:80]}")

    if APPLY_C5:
        # Migicovsky 等人间空格
        no_space = full.count("Migicovsky等人")
        with_space = full.count("Migicovsky 等人")
        print(f"  ✓ C5 Migicovsky 空格 (无空格残留: {no_space}, 有空格: {with_space})")

    # 残留分号检测
    print()
    print("剩余英文分号检测（仅在引用括号内）：")
    leftovers = []
    for p in doc2.paragraphs:
        t = p.text
        if is_reference_para(t):
            continue
        for m in re.finditer(r"（[^（）()]*?;[^（）()]*?）", t):
            leftovers.append(m.group(0)[:80])
    if leftovers:
        print(f"  ⚠️ 还有 {len(leftovers)} 处")
        for s in leftovers:
            print(f"    {s}")
    else:
        print("  ✓ 无残留")


if __name__ == "__main__":
    main()
